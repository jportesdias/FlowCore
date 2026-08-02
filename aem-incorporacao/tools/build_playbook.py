from pathlib import Path
from PIL import Image
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION_START
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(r"F:\.FlowCore Solutions\Site FlowCore Solutions\60_PROJETOS\A&M incorporação")
ASSETS = ROOT / "work" / "manual_pages"
OUT = ROOT / "output"
OUT.mkdir(exist_ok=True)
OUTFILE = OUT / "CAL-1211_Procedimento_e_Anexo_Checklist_Agilent_490_Micro_GC.docx"

INK = "1E2933"
MUTED = "66717D"
ACCENT = "244B5A"
LINE = "CBD2D8"
LIGHT = "F2F4F5"
WARN = "8A3B32"
WHITE = "FFFFFF"

def crop(src, dst, box):
    im = Image.open(src)
    im.crop(box).save(dst)

crop(ASSETS / "manual_p17.png", ASSETS / "fig_front.png", (90, 190, 930, 1085))
crop(ASSETS / "manual_p18.png", ASSETS / "fig_back.png", (85, 155, 1005, 790))
crop(ASSETS / "manual_p46.png", ASSETS / "fig_filter.png", (85, 360, 1010, 1165))

doc = Document()
sec = doc.sections[0]
sec.top_margin = Inches(.62); sec.bottom_margin = Inches(.58)
sec.left_margin = Inches(.66); sec.right_margin = Inches(.66)
sec.header_distance = Inches(.28); sec.footer_distance = Inches(.28)

styles = doc.styles
normal = styles['Normal']
normal.font.name = 'Arial'; normal.font.size = Pt(9); normal.font.color.rgb = RGBColor.from_string(INK)
normal.paragraph_format.space_after = Pt(3); normal.paragraph_format.line_spacing = 1.08
for name, size, before, after, color in [
    ('Title', 25, 0, 8, INK), ('Subtitle', 11, 0, 10, MUTED),
    ('Heading 1', 15, 8, 5, ACCENT), ('Heading 2', 11, 6, 3, INK), ('Heading 3', 9.5, 4, 2, ACCENT)]:
    s = styles[name]; s.font.name='Arial'; s.font.size=Pt(size); s.font.color.rgb=RGBColor.from_string(color)
    s.font.bold = name != 'Subtitle'; s.paragraph_format.space_before=Pt(before); s.paragraph_format.space_after=Pt(after)
    s.paragraph_format.keep_with_next = True

def shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr(); shd = tcPr.find(qn('w:shd'))
    if shd is None: shd = OxmlElement('w:shd'); tcPr.append(shd)
    shd.set(qn('w:fill'), fill)

def borders(table, color=LINE, size='4'):
    tblPr = table._tbl.tblPr; el = tblPr.find(qn('w:tblBorders'))
    if el is None: el = OxmlElement('w:tblBorders'); tblPr.append(el)
    for edge in ('top','left','bottom','right','insideH','insideV'):
        x = OxmlElement(f'w:{edge}'); x.set(qn('w:val'),'single'); x.set(qn('w:sz'),size); x.set(qn('w:color'),color); el.append(x)

def set_cell_margins(cell, top=60, start=80, bottom=60, end=80):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr(); tcMar = tcPr.first_child_found_in('w:tcMar')
    if tcMar is None: tcMar=OxmlElement('w:tcMar'); tcPr.append(tcMar)
    for m,v in [('top',top),('start',start),('bottom',bottom),('end',end)]:
        node=tcMar.find(qn(f'w:{m}'))
        if node is None: node=OxmlElement(f'w:{m}'); tcMar.append(node)
        node.set(qn('w:w'),str(v)); node.set(qn('w:type'),'dxa')

def table(headers, rows, widths=None, font=7.5, header_fill=ACCENT):
    t=doc.add_table(rows=1, cols=len(headers)); t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.autofit=False
    for i,h in enumerate(headers):
        c=t.rows[0].cells[i]; c.text=h; shade(c,header_fill); set_cell_margins(c)
        for r in c.paragraphs[0].runs: r.font.name='Arial'; r.font.size=Pt(font); r.font.bold=True; r.font.color.rgb=RGBColor.from_string(WHITE)
    for row in rows:
        cells=t.add_row().cells
        for i,val in enumerate(row):
            cells[i].text=str(val); set_cell_margins(cells[i]); cells[i].vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cells[i].paragraphs:
                p.paragraph_format.space_after=Pt(0); p.paragraph_format.line_spacing=1.0
                for r in p.runs: r.font.name='Arial'; r.font.size=Pt(font); r.font.color.rgb=RGBColor.from_string(INK)
    if widths:
        for row in t.rows:
            for i,w in enumerate(widths): row.cells[i].width=Inches(w)
        grid=t._tbl.tblGrid
        for child in list(grid): grid.remove(child)
        for w in widths:
            gc=OxmlElement('w:gridCol'); gc.set(qn('w:w'),str(int(w*1440))); grid.append(gc)
        tblW=t._tbl.tblPr.find(qn('w:tblW'))
        if tblW is None: tblW=OxmlElement('w:tblW'); t._tbl.tblPr.append(tblW)
        tblW.set(qn('w:w'),str(int(sum(widths)*1440))); tblW.set(qn('w:type'),'dxa')
    borders(t)
    return t

def p(text='', bold=False, color=INK, size=9, align=None, after=3, keep=False):
    para=doc.add_paragraph(); para.paragraph_format.space_after=Pt(after); para.paragraph_format.keep_with_next=keep
    if align is not None: para.alignment=align
    r=para.add_run(text); r.font.name='Arial'; r.font.size=Pt(size); r.font.bold=bold; r.font.color.rgb=RGBColor.from_string(color)
    return para

def bullets(items, size=8.5):
    for item in items:
        para=doc.add_paragraph(style='List Bullet'); para.paragraph_format.left_indent=Inches(.2); para.paragraph_format.first_line_indent=Inches(-.14); para.paragraph_format.space_after=Pt(1.5)
        r=para.add_run(item); r.font.name='Arial'; r.font.size=Pt(size)

def box(title, text, fill=LIGHT, title_color=ACCENT):
    t=doc.add_table(rows=1, cols=1); t.autofit=False; t.columns[0].width=Inches(7.02); c=t.cell(0,0); shade(c,fill); set_cell_margins(c,100,140,100,140); borders(t,color=fill,size='2')
    q=c.paragraphs[0]; q.paragraph_format.space_after=Pt(2); a=q.add_run(title); a.bold=True; a.font.name='Arial'; a.font.size=Pt(9); a.font.color.rgb=RGBColor.from_string(title_color)
    q2=c.add_paragraph(); q2.paragraph_format.space_after=Pt(0); b=q2.add_run(text); b.font.name='Arial'; b.font.size=Pt(8.5); b.font.color.rgb=RGBColor.from_string(INK)

def newpage(): doc.add_page_break()

def page_header_footer():
    for section in doc.sections:
        hp=section.header.paragraphs[0]; hp.text='PLAYBOOK DE CAMPO  |  CAL-1211'; hp.alignment=WD_ALIGN_PARAGRAPH.RIGHT
        for r in hp.runs: r.font.name='Arial'; r.font.size=Pt(7); r.font.color.rgb=RGBColor.from_string(MUTED)
        fp=section.footer.paragraphs[0]; fp.alignment=WD_ALIGN_PARAGRAPH.RIGHT
        r=fp.add_run('Documento de trabalho  |  '); r.font.name='Arial'; r.font.size=Pt(7); r.font.color.rgb=RGBColor.from_string(MUTED)
        fld=OxmlElement('w:fldSimple'); fld.set(qn('w:instr'),'PAGE'); fp._p.append(fld)

# 1 — abertura
p('PLAYBOOK TÉCNICO DE EXECUÇÃO', bold=True, color=ACCENT, size=9, after=8)
doc.add_paragraph('Calibração de analisador\ncromatográfico on-line', style='Title')
p('Agilent 490 Micro GC  |  CAL-1211', color=MUTED, size=11, after=18)
table(['Campo','Definição'],[
('Objetivo','Executar e documentar a verificação/calibração do analisador com gás certificado, preservando a configuração e a rastreabilidade.'),
('Escopo','Inspeção, baseline, conexão do gás, estabilização, corridas, avaliação, registro e retorno à operação. Não inclui reparo interno.'),
('Condição operacional','A confirmar com Operação: medição congelada/transferida, fluxo de calibração liberado e intervenção autorizada.'),
('Equipe de referência','1 técnico de medição + 1 técnico de laboratório (CAL-1211).'),
('Duração','3 h de mão de obra (CAL-1211). A estabilização indicada no documento-base é de no mínimo 6 h; planejar a janela total.'),
('Local / tag / nº de série','A preencher antes da execução.'),
], widths=[1.55,5.47],font=8.2)
doc.add_heading('Resultado final esperado', level=1)
bullets(['Analisador disponível, sem falhas críticas ou fatais ativas e com comunicação estável.', 'Resultados avaliados contra o método aprovado e os requisitos de confirmação metrológica.', 'Configuração original preservada ou alteração formalmente aprovada e registrada.', 'CROM-M-01, certificado do gás, dados brutos, cromatogramas e registro da ordem de serviço anexados.'])
box('REGRA DE LIBERAÇÃO','Não iniciar sem identificar positivamente equipamento, método, gás de calibração, concentrações certificadas e critério de aceitação vigente.')

# 2 — papéis, pré-requisitos, riscos
newpage(); p('ANEXO A', bold=True, color=ACCENT, size=9, after=3)
doc.add_heading('Checklist de execução de campo', level=1)
box('INSTRUÇÃO DE USO','Preencher este anexo durante a atividade. Marcar cada item somente após confirmar o resultado e registrar a evidência correspondente. Itens não aplicáveis exigem justificativa.')
doc.add_heading('A.1 Liberação e controle do trabalho', level=2)
table(['Função','Responsabilidade'],[
('Execução técnica','Executar etapas, registrar baseline, dados e desvios.'),('Laboratório','Confirmar certificado, mistura, validade e rastreabilidade do gás.'),('Operação / liberação','Autorizar intervenção, congelamento/transferência e retorno ao serviço.'),('Supervisão técnica','Aprovar método, critérios, desvios e alterações.'),('Encerramento','Validar registros e fechar PT/WO conforme sistema local.')],widths=[1.55,5.47],font=8)
doc.add_heading('Pré-requisitos', level=2)
table(['Verificação','Critério para prosseguir','Resp.','Status'],[
('PT e análise de risco','Emitidas, comunicadas e barreiras implantadas.','Liberação','☐'),('Tag e configuração','Equipamento e canais confirmados no campo e no CDS.','Execução','☐'),('Impacto da medição','Congelamento/transferência autorizado e partes informadas.','Operação','☐'),('Gás de calibração','Certificado válido; componentes/faixas compatíveis com o método.','Laboratório','☐'),('Gás de arraste','Tipo corresponde à configuração; suprimento seco e limpo.','Execução','☐'),('Instrumentos de teste','Identificados, adequados e com calibração válida.','Execução','☐'),('Backup','Método, fatores, integração e configuração exportados.','Execução','☐')],widths=[1.5,3.55,1.25,.72],font=7.1)
doc.add_heading('Riscos críticos e barreiras', level=2)
table(['Risco','Barreira obrigatória','Parar quando'],[
('Gás inflamável / pressurizado','Ventilação, teste de vazamento, conexões compatíveis e controle da fonte.','Vazamento, atmosfera insegura ou pressão não controlada.'),('Superfície quente','Confirmar temperatura antes de tocar linhas aquecidas.','Temperatura incompatível com a intervenção.'),('Perda de medição','Autorização e comunicação contínua com Operação.','Comunicação perdida ou impacto não previsto.'),('Alteração indevida','Backup e dupla verificação antes de aplicar método/fator.','Valor original ausente ou alteração não autorizada.')],widths=[1.65,3.4,1.97],font=7.2)

# 3 — recursos e baseline
newpage(); doc.add_heading('A.2 Recursos e baseline', level=1)
table(['Item','Especificação mínima','Finalidade / verificação'],[
('Laptop + CDS','Software e licença compatíveis; cabo/rede disponível.','Conectar, exportar backup e executar sequência.'),('Gás de calibração','Mistura certificada, dentro da validade e compatível com o método.','Conferir certificado, lote, pressão e concentrações.'),('Gás de arraste','He/H₂/N₂/Ar conforme configuração; pureza mín. 99,999%.','Pressão de entrada 550 ± 10 kPa (80 ± 1,5 psi), manual p. 21/33.'),('Regulador e linhas','Regulador de dois estágios; tubing metálico limpo para GC.','Não usar tubing plástico; testar vazamento.'),('Filtro externo','Elemento de 5 µm e conexões adequadas.','Conferir sentido, condição e ausência de umidade.'),('Registros','CROM-M-01, WO e meios para fotos/PDF.','Garantir rastreabilidade antes, durante e depois.')],widths=[1.3,2.65,3.07],font=7.4)
doc.add_heading('Registro da condição inicial', level=2)
bullets(['Fotografar equipamento, tag, cilindros, reguladores, conexões e posição das válvulas.', 'Registrar nº de série, canais, método ativo, versão de software/firmware, IP e estado da comunicação.', 'Exportar configuração, fatores de resposta, tabela de componentes, integração e último cromatograma válido.', 'Registrar LEDs, alarmes, pressões, temperaturas, fluxo selecionado e valores vistos no CDS/sistema de controle.', 'Inspecionar cabos, tubing, filtros, aterramento, ventilação e sinais de vazamento/condensação.'])
box('BASELINE CONFIRMADA','☐ Identificação positiva  ☐ Backup legível  ☐ Condição registrada  ☐ Impacto operacional controlado  ☐ Autorização para prosseguir')
table(['Vista frontal','Vista traseira'],[('', '')],widths=[3.51,3.51],font=7.2)
t=doc.tables[-1]
for cell,img in zip(t.rows[1].cells,[ASSETS/'fig_front.png',ASSETS/'fig_back.png']):
    cell.text=''; q=cell.paragraphs[0]; q.alignment=WD_ALIGN_PARAGRAPH.CENTER; q.add_run().add_picture(str(img),width=Inches(3.1))
p('Figuras adaptadas do Agilent 490 Micro GC User Manual, p. 17–18.',color=MUTED,size=7,after=0)

# 4 — execução I
newpage(); doc.add_heading('A.3 Execução passo a passo', level=1)
table(['#','Ação de campo','Como executar','Resultado / evidência','Parar quando','Status'],[
('1','Confirmar liberação','Revisar PT, risco, tag, impacto e comunicação.','Autorização registrada.','Barreira ausente ou escopo divergente.','☐'),
('2','Validar o gás certificado','Conferir lote, validade, componentes, incertezas e concentrações contra o método.','Certificado anexado e compatibilidade confirmada.','Componente/faixa incompatível ou certificado inválido.','☐'),
('3','Confirmar o gás de arraste','Comparar tipo físico com configuração do CDS. Confirmar 550 ± 10 kPa e pureza mín. 99,999%.','Tipo, pressão e pureza registrados.','Tipo divergente, vazamento ou pressão instável.','☐'),
('4','Inspecionar a linha de amostra','Confirmar linha aberta/desobstruída, filtro correto, ausência de líquido e vent seguro.','Foto e inspeção registradas.','Obstrução, umidade/condensação ou descarga insegura.','☐'),
('5','Registrar e exportar o baseline','Salvar método, fatores, integração, alarmes e cromatograma anterior.','Pacote de backup recuperável.','Backup incompleto ou ilegível.','☐'),
('6','Conectar o laptop','Usar LAN/cabo autorizado; abrir a sessão on-line e confirmar instrumento/canais.','Comunicação estável e equipamento correto.','Conexão intermitente ou instrumento incorreto.','☐'),
('7','Selecionar o fluxo de calibração','Com Operação ciente, posicionar a seleção conforme P&ID/método aprovado.','Fluxo correto confirmado por dois meios.','Posição duvidosa ou impacto não autorizado.','☐'),
('8','Estabilizar o sistema','Aplicar o método aprovado. Aguardar no mínimo 6 h conforme CAL-1211 e até o estado Ready/condições estáveis.','Tendências e horários de início/fim registrados.','Erro crítico/fatal ou instabilidade sem causa conhecida.','☐'),
],widths=[.28,1.16,2.3,1.55,1.25,.48],font=6.4)
box('HOLD POINT HP-01','Antes de executar corridas de calibração: confirmar baseline, certificado, fluxo correto, estabilidade e autorização da supervisão técnica.')

# 5 — execução II e figura
newpage(); doc.add_heading('A.3 Execução passo a passo — continuação', level=1)
table(['#','Ação de campo','Como executar','Resultado / evidência','Parar quando','Status'],[
('9','Executar a sequência','Executar a sequência de calibração definida no método aprovado. Quando aplicável ao checkout, usar ao menos 3 corridas (manual p. 41).','Dados brutos e cromatogramas salvos.','Corrida abortada, pico não identificado ou condição fora do método.','☐'),
('10','Avaliar o cromatograma','Verificar baseline, retenção, integração, identificação e separação dos picos contra o método/histórico.','Avaliação por canal registrada.','Coeluição indevida, ruído/deriva ou identificação duvidosa.','☐'),
('11','Calcular o desempenho','Comparar resultado com valor certificado e calcular erro/repetibilidade pelo requisito metrológico vigente.','Planilha/relatório rastreável.','Critério ausente ou resultado reprovado.','☐'),
('12','Decidir sobre ajuste','Se aprovado, não alterar fatores. Se reprovado, diagnosticar antes de ajustar. Registrar valor original e obter autorização.','Decisão e justificativa registradas.','Causa desconhecida ou ajuste não autorizado.','☐'),
('13','Aplicar ajuste autorizado','Alterar somente o parâmetro aprovado; salvar versão anterior e nova.','Rastreabilidade antes/depois.','Alteração irreversível sem backup.','☐'),
('14','Repetir e validar','Executar nova sequência e reavaliar todos os canais/componentes afetados.','Resultado final dentro do critério aprovado.','Falha persistente ou efeito adverso em outro componente.','☐'),
('15','Restaurar a operação','Retornar fluxo, modo e interfaces conforme alinhamento com Operação. Confirmar valores coerentes.','Aceite da Operação registrado.','Alarme ativo, comunicação instável ou valor incoerente.','☐'),
('16','Encerrar e anexar registros','Concluir CROM-M-01; anexar certificado, cromatogramas e PDF da calibração ao WO/AMOS.','Pacote final completo e pendências abertas.','Evidência ausente ou desvio sem tratamento.','☐')],widths=[.28,1.16,2.3,1.55,1.25,.48],font=6.4)
doc.add_heading('Conexão do filtro externo', level=2)
q=doc.add_paragraph(); q.alignment=WD_ALIGN_PARAGRAPH.CENTER; q.paragraph_format.space_after=Pt(1); q.add_run().add_picture(str(ASSETS/'fig_filter.png'),width=Inches(4.5))
p('No conjunto do manual: apertar manualmente e depois mais 1/8 de volta com chave 7/16 pol.; orientar a seta da parte fêmea para a conexão do Micro GC. Confirmar aplicabilidade à instalação local. Fonte: manual, p. 46.',color=MUTED,size=7,after=0)

# 6 — hold, critérios, stop
newpage(); doc.add_heading('A.4 Hold points e aceitação', level=1)
table(['Hold point','Condição','Evidência','Autorização'],[
('HP-01','Baseline, gases, fluxo e estabilidade confirmados.','Checklist + certificado + backup.','Supervisão técnica'),('HP-02','Qualquer ajuste de fator, integração ou método.','Resultado reprovado + diagnóstico + valor original.','Aprovador de alteração'),('HP-03','Retorno ao fluxo de processo/automático.','Resultados finais + ausência de falha crítica.','Operação'),('HP-04','Encerramento da atividade.','CROM-M-01 e pacote de registros completo.','Responsável pelo encerramento')],widths=[.75,2.55,2.35,1.37],font=7.3)
doc.add_heading('Critérios de aceitação', level=2)
table(['Item','Critério','Método / registro'],[
('Gás de calibração','Certificado válido e mistura compatível com o método.','Certificado e conferência do laboratório.'),('Condições do instrumento','Estado Ready e sem erro classe 2/3.','Status do CDS, LEDs e log de erros.'),('Cromatograma','Picos identificados, integrados e separados conforme método aprovado/histórico aceito.','Cromatogramas por canal.'),('Erro e repetibilidade','Critério a confirmar antes do início da atividade no Guia Técnico de Medição/método vigente.','Cálculo e relatório de confirmação metrológica.'),('Configuração','Original restaurada ou alteração formalmente aprovada e rastreada.','Comparativo de backup.'),('Retorno','Comunicação estável, valores coerentes e aceite da Operação.','Tendência, tela e assinatura/registro.')],widths=[1.35,3.65,2.02],font=7.4)
box('STOP WORK','Interromper imediatamente diante de identificação duvidosa; condição diferente da prevista; pressão/temperatura/energia não controlada; atmosfera insegura; vazamento; perda de comunicação; divergência documental; instrumento de teste inadequado; resultado fora do limite sem causa conhecida; mudança de escopo; risco de dano; improvisação; ou ausência de responsável autorizado. Retomar somente após avaliar, registrar e autorizar formalmente.',fill='F7ECEA',title_color=WARN)
doc.add_heading('Referência rápida de falhas', level=2)
table(['Indicação','Interpretação / resposta'],[
('Ready apagado','Sistema não pronto: verificar status do CDS, condições do método e estabilização.'),('Error piscando','Erro presente: consultar log e classificação. Não aceitar com classe 2/3.'),('Pressão baixa','Manual: abaixo de 35 kPa gera aviso; verificar suprimento. Não confundir com pressão de entrada especificada.'),('Falha de calibração do TCD','Erro classe 2 no manual; interromper e encaminhar diagnóstico autorizado.')],widths=[1.55,5.47],font=7.5)

# 7 — diagnóstico, retorno, registros
newpage(); doc.add_heading('A.5 Diagnóstico e retorno à operação', level=1)
table(['Sintoma','Verificar primeiro','Se normal','Se anormal / bloqueio'],[
('Resultado deslocado','Certificado, unidade, componente, fluxo selecionado e integração.','Comparar retenção/fator com backup.','Corrigir somente identificação/configuração autorizada; bloquear se causa incerta.'),('Baixa repetibilidade','Pressão, vazamento, umidade, filtro, estabilidade e linha desobstruída.','Repetir sequência controlada.','Eliminar causa externa; não ajustar fator para mascarar instabilidade.'),('Pico ausente/coeluído','Mistura, canal, tabela de componentes, retenção e método.','Comparar com histórico e checkout.','Bloquear ajuste sem diagnóstico e aprovação.'),('Sem comunicação','Energia, LEDs, cabo, IP/sub-rede e sessão CDS.','Reconectar sem alterar configuração.','Bloquear reset de IP/reinicialização sem autorização.'),('Error ativo','Classe, código, canal e hora no log.','Tratar conforme manual e repetir verificação.','Classe 2 para a corrida; classe 3 com shutdown: não prosseguir.')],widths=[1.25,2.4,1.55,1.82],font=7.1)
doc.add_heading('Checklist de retorno', level=2)
table(['☐','Confirmação','☐','Confirmação'],[
('☐','Ferramentas e materiais removidos.','☐','Tubing, tampas e conexões reinstalados.'),('☐','Válvulas na posição autorizada.','☐','Isolamentos removidos conforme liberação.'),('☐','Parâmetros restaurados/aprovados.','☐','Alarmes analisados e normalizados.'),('☐','Comunicação estável.','☐','Valores coerentes em campo/controle.'),('☐','Operação informada e aceite registrado.','☐','Área limpa; pendências registradas.'),('☐','Evidências anexadas.','☐','Responsáveis e horários identificados.')],widths=[.35,3.16,.35,3.16],font=7.4)
box('SISTEMA APTO PARA RETORNO À OPERAÇÃO','☐ Critérios atendidos  ☐ HP-03 liberado  ☐ Operação aceitou  ☐ Condição final registrada')
doc.add_heading('Registros obrigatórios', level=2)
bullets(['Fotos antes/depois; tag, cilindros, reguladores e conexões.', 'Certificado do gás e dos instrumentos utilizados.', 'Valores certificados, encontrados, aplicados e resultados calculados.', 'Cromatogramas e dados brutos de todas as corridas utilizadas.', 'Configurações e fatores originais/finais; alarmes e logs.', 'Desvios, diagnóstico, ações autorizadas, peças e pendências.', 'Pessoas, data/horário, condição operacional, documentos consultados e aceite final.'],size=8)

# 8 — lacunas, fechamento e fontes
newpage(); doc.add_heading('A.6 Pontos a confirmar antes da emissão', level=1)
table(['Lacuna','Fonte / responsável para confirmação','Bloqueia'],[
('Tag, local, configuração e nº de série do Micro GC.','Cadastro, folha de dados e inspeção de campo.','Execução'),('Método aprovado, canais, componentes e tempos de retenção.','CDS, IOM específica e responsável técnico.','Corridas'),('Mistura, concentrações e incerteza do gás.','Certificado vigente e Laboratório.','Execução'),('Critério de erro, repetibilidade e regra de ajuste.','Guia Técnico de Medição / confirmação metrológica.','Aceitação'),('Pressão dos gases de calibração/amostra e limites de troca de cilindro.','Projeto, datasheet/IOM e responsável técnico.','Conexão'),('P&ID, posição de válvulas e destino dos vents.','Documentação as-built e Operação.','Intervenção'),('Estratégia de congelamento/transferência da medição.','Procedimento operacional e Operação.','Liberação'),('Compatibilidade com área classificada e requisitos de EPI.','Classificação de área, PT e análise de risco.','Acesso'),('Software, credenciais, IP e permissões para backup/ajuste.','Administração do CDS / automação.','Conexão/ajuste')],widths=[3.0,3.25,.77],font=7.2)
doc.add_heading('Encerramento executivo', level=2)
table(['Campo','Registro final'],[
('Atividade executada','____________________________________________________________'),('Resultado obtido','____________________________________________________________'),('Critérios atendidos','☐ Sim  ☐ Não  ☐ Parcial  |  Evidência: __________________________'),('Desvios pendentes','____________________________________________________________'),('Condição atual','☐ Em operação  ☐ Indisponível  ☐ Operação restrita'),('Ação futura / responsável','____________________________________________________________')],widths=[1.6,5.42],font=8)
doc.add_heading('Referências utilizadas', level=2)
bullets(['CAL-1211 (PT-BR) — Calibrar analisador cromatógrafo gasoso on-line.', 'Agilent 490 Micro GC User Manual, G3581-90001, 6ª ed., 2017.', 'Agilent 490 Micro GC Solution Data Sheet, 5991-6034EN, 2017.', 'CROM-M-01 — History Template; Guia Técnico de Medição e documentos locais citados no CAL-1211 (a disponibilizar).'],size=7.7)
p('Nota de controle: este playbook consolida as fontes fornecidas. Dados dependentes da unidade permanecem marcados para confirmação antes da emissão/aprovação.',color=MUTED,size=7.5,after=0)

page_header_footer()
doc.core_properties.title='Playbook de Campo - Calibração Agilent 490 Micro GC'
doc.core_properties.subject='CAL-1211'
doc.core_properties.author=''
doc.save(OUTFILE)
print(OUTFILE)
